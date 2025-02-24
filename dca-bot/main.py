import os
import pandas as pd
from zipfile import ZipFile
import shutil
from pathlib import Path
import credentials
import base64
# from skpy import Skype
from skpy import SkypeEventLoop, SkypeNewMessageEvent
import pyodbc
import paramiko
import re
from PIL import Image, ImageDraw, ImageFont
import stat
from docx import Document
from docx2pdf import convert
import glob
import zipfile

# -------------------------------------------------------------------

ssh = paramiko.SSHClient()
ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())


# -------------------------------------------------------------------


class SkypeListener(SkypeEventLoop):
    """ Listen to a channel continuously """

    def __init__(self):
        username = credentials.skypeUser
        password = base64.b64decode(credentials.skypePassword).decode("utf-8")
        token_file = '.tokens-app'
        super(SkypeListener, self).__init__(username, password, token_file)


        self.LOANEID_pattern = re.compile(r'((mir)|(vcr)|(ocr)|(atm)|(sen)|(mnc)|(dcr)|(mnv)|(shi)|(tkm))\d+')

        #self.VID_Something_chatId = '19:e940d0a9489a4c2b8e3a17f85fc234ba@thread.skype' #prod
        self.VID_Something_chatId ='19:03fb1a3c75384a6db49f4c02c4bf4414@thread.skype' #test
        self.VID_Something_Group = self.chats[self.VID_Something_chatId]
        self.GetInfo_pattern = re.compile(
            r'^getinfo\s+((mir)|(vcr)|(ocr)|(atm)|(sen)|(mnc)|(dcr)|(mnv)|(shi)|(tkm))\d+\s*$')
        # self.GetInfo_pattern = re.compile(r'^getinfo\s+((mir)|(atm)|(sen))\d+\s*$')


        self.VID_Something_Group.sendMsg("I'm listening !")


        self.SMS_Zalo_Face_chatId = '19:03fb1a3c75384a6db49f4c02c4bf4414@thread.skype'
        self.SMS_Zalo_Face_Group = self.chats[self.SMS_Zalo_Face_chatId]
        self.GetSMS_pattern = re.compile(
            r'^getsms[123]\s+((mir)|(vcr)|(ocr)|(atm)|(sen)|(mnc)|(dcr)|(mnv)|(shi)|(tkm))\d+\s*$')
        self.SMStemplate_pattern = re.compile('^getsms[123]')

        self.VID_DataUpdate_chatId = '19:03fb1a3c75384a6db49f4c02c4bf4414@thread.skype'
        self.VID_DataUpdate_Group = self.chats[self.VID_DataUpdate_chatId]
        self.GetData_pattern = re.compile(
            r'^getdata\s+((mir)|(vcr)|(ocr)|(atm)|(sen)|(mnc)|(dcr)|(mnv)|(shi)|(tkm))\d+\s*$')

        self.VID_DK_chatId = '19:03fb1a3c75384a6db49f4c02c4bf4414@thread.skype'
        self.VID_DK_Group = self.chats[self.VID_DK_chatId]
        self.GetData_pattern = re.compile(
            r'^dk\s+((mir)|(vcr)|(ocr)|(atm)|(sen)|(mnc)|(dcr)|(mnv)|(shi)|(tkm))\d+\s*$')
        self.Legal_pattenrn = re.compile(
            r'^legal\s+((mir)|(vcr)|(ocr)|(atm)|(sen)|(mnc)|(dcr)|(mnv)|(shi)|(tkm))\d+\s*$')

        self.VID_TBKK_sunshine = '19:10ea21d96d764d549b4f409884548518@thread.skype'
        #nho doi ten
        self.VID_TBKK_concentrix = '19:605dc4e4be8b49e2a6effe6233f386fe@thread.skype'
        #self.VID_TBKK_Group = self.chats[self.VID_TBKK_chatId]
        self.TBKK_pattern = re.compile(
            r'^tbkk\s+((mir)|(vcr)|(ocr)|(atm)|(sen)|(mnc)|(dcr)|(mnv)|(shi)|(tkm))\d+\s*$')

        self.SMS_font = ImageFont.truetype("Tahoma.ttf", 14)
        self.SMS_Wartermark_font = ImageFont.truetype("Tahoma.ttf", 34)

    def onEvent(self, event):
        # for request in self.contacts.requests():
        #     request.accept()

        if isinstance(event, SkypeNewMessageEvent):
            message = {"user_id": event.msg.userId,
                       "chat_id": event.msg.chatId,
                       "msg": event.msg.content}
            print(message["chat_id"])
            if message["chat_id"] == self.VID_Something_chatId:
                # if message["chat_id"] == "abc":
                print(message)
                if self.GetInfo_pattern.match(message["msg"].lower()) != None:

                    SearchValue = self.LOANEID_pattern.search(message["msg"].lower()).group()
                    if SearchValue != None:
                        myDBconn = pyodbc.connect(credentials.myLocalDBConn)

                        # print('\nRequested:', message["user_id"], '-', message["chat_id"], '-', message["msg"])
                        self.VID_Something_Group.sendMsg(
                            '@{}: Yêu cầu của bạn đang được xử lí'.format(message["user_id"]))

                        result = pd.read_sql_query("""
                                EXEC [localDB].[dbo].[SkypeGetInfo]
                                @userName = '{}'
                                ,@Domain = '{}'
                                ,@SearchKey = {}
                                ,@SearchValue = '{}'
                            """.format(message["user_id"], 'Inhouse', 1, SearchValue)
                                                   , myDBconn)

                        myDBconn.commit()
                        myDBconn.close()

                        rowCount = result['ID'].count()

                        if rowCount == 0:
                            # print('Status: "{}" not found'.format(SearchValue))
                            self.VID_Something_Group.sendMsg(
                                '@{}: Không tìm thấy hợp đồng "{}"'.format(message["user_id"], SearchValue))
                        else:
                            # print('Path: {}'.format(result["Path"][0]))

                            destinationPath = os.path.join(r"Output", SearchValue)

                            if os.path.exists(destinationPath):
                                shutil.rmtree(destinationPath)

                            Path(destinationPath).mkdir(
                                parents=True
                                # ,exist_ok=True
                            )

                            ssh.connect(
                                credentials.sftpHost,
                                credentials.sftpPort,
                                "nguyen.dao",
                                "rYCtjDzpJ5Weh5O"
                                # credentials.sftpUser,
                                # base64.b64decode(credentials.sftpPassword).decode("utf-8")
                            )
                            sftp = ssh.open_sftp()

                            #             # shutil.copy(sourcePath, destinationPath)
                            for path in result["Path"]:
                                temp = sftp.lstat(path)
                                if stat.S_ISDIR(temp.st_mode):
                                    for file in sftp.listdir(path):
                                        # shutil.copy(os.path.join(result["Path"], file), destinationPath)
                                        fullSourcePath = path + '/' + file
                                        fullDestPath = destinationPath + '\\' + file
                                        sftp.get(fullSourcePath, fullDestPath)
                                        print(fullSourcePath, '-', fullDestPath)
                                else:
                                    fullSourcePath = path
                                    fullDestPath = destinationPath + '\\' + os.path.basename(path)
                                    sftp.get(fullSourcePath, fullDestPath)
                                    print(fullSourcePath, '-', fullDestPath)

                            ssh.close

                            with ZipFile(os.path.join(r".\Output", SearchValue + '.zip'), 'w') as zipObj:
                                for folderName, subfolders, filenames in os.walk(destinationPath):
                                    for filename in filenames:
                                        # create complete filepath of file in directory
                                        filePath = os.path.join(folderName, filename)
                                        # Add file to zip
                                        zipObj.write(filePath, os.path.basename(filePath))

                            zipObj.close()

                            fullPath = os.path.join(r".\Output", SearchValue + '.zip')
                            fileName = os.path.basename(fullPath)

                            try:
                                sent2User = self.contacts[message["user_id"]].chat

                                # sent2User = sk.contacts[msg.userId].chat
                                # sent2User.sendMsg('test Msg')

                                sent2User.sendFile(open(fullPath, "rb"), fileName, image=False)  # file upload
                                self.VID_Something_Group.sendMsg(
                                    '@{}: Hợp đồng "{}" đã được gửi cho bạn'.format(message["user_id"], SearchValue))

                            except ValueError:
                                print("Oops! Check your contacts & try again...")

                elif self.Legal_pattenrn.match(message["msg"].lower()) != None:
                    SearchValue = self.LOANEID_pattern.search(message["msg"].lower()).group()
                    Skype_legal = ["live:.cid.f1384fdd9713f8f1", "live:tunguyen24197", "live:.cid.118e632229c91974",
                                   "live:.cid.c6d02e974b5fc509", "live:.cid.c012fb30e14c9af5"]
                    if SearchValue != None and message["user_id"] in Skype_legal:
                        self.VID_Something_Group.sendMsg(
                            '@{}: Yêu cầu của bạn đang được xử lí'.format(message["user_id"]))

                        def zip_directory(directory_path, zip_path):
                            if not os.path.exists(directory_path):
                                print("Lỗi: Thư mục không tồn tại")
                                return

                            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                                for root, dirs, files in os.walk(directory_path):
                                    for file in files:
                                        file_path = os.path.join(root, file)
                                        zipf.write(file_path, os.path.relpath(file_path, directory_path))

                            print("Thư mục đã được nén thành công")

                        # Thay đổi đường dẫn thư mục và đường dẫn nén tùy theo yêu cầu của bạn
                        port = SearchValue[:3]
                        directory_path = "C:/Users/Admin/Downloads/25. 300 case gửi CNX/" + port + "/" + SearchValue
                        zip_path = "C:/Users/Admin/Downloads/zip/" + SearchValue + ".zip"
                        zip_directory(directory_path, zip_path)

                        if os.path.exists(directory_path):
                            try:
                                sent2User = self.contacts[message["user_id"]].chat

                                # sent2User = sk.contacts[msg.userId].chat
                                # sent2User.sendMsg('test Msg')

                                fullPath = os.path.join("C:/Users/Admin/Downloads/zip", SearchValue + '.zip')
                                fileName = os.path.basename(fullPath)
                                sent2User.sendFile(open(fullPath, "rb"), fileName, image=False)  # file upload
                                self.VID_Something_Group.sendMsg(
                                    '@{}: Hợp đồng "{}" đã được gửi cho bạn'.format(message["user_id"], SearchValue))

                            except ValueError:
                                print("Oops! Check your contacts & try again...")
                        else:
                            self.VID_Something_Group.sendMsg(
                                '@{}: Không tìm thấy hợp đồng'.format(message["user_id"]))
                            
                elif message["msg"].lower() == "help":
                    self.VID_Something_Group.sendMsg(
                        "Xin vui lòng xem một số chức năng bên dưới:\n"
                        "1. Lấy thông tin, cú pháp 'getinfo + LOANIED'\n"
                        "2. Kiểm tra trạng thái BOT, gửi 'hello' hoặc 'hi' nếu được phản hồi thì trạng thái vẫn đang hoạt động.\n"
                        "3. Truy cập trang quản lý http://bot2.han.vnfin.vn"
                    )
                elif message["msg"].lower() == "hi" or message["msg"].lower() == "hello":
                    self.VID_Something_Group.sendMsg("Xin chào, Tôi là VID BOT. Xin vui lòng gửi yêu cầu !!")

            elif message["chat_id"] == self.SMS_Zalo_Face_chatId:
                if self.GetSMS_pattern.match(message["msg"].lower()) != None:
                    SearchValue = self.LOANEID_pattern.search(message["msg"].lower()).group()
                    Template = self.SMStemplate_pattern.search(message["msg"].lower()).group()[-1]
                    if SearchValue != None:
                        # print(SearchValue)
                        # print(message.userId)
                        # print(message.chatId)
                        # print(message.userId)
                        self.SMS_Zalo_Face_Group.sendMsg(
                            '@{}: Yêu cầu của bạn đang được xử lí'.format(message["user_id"]))

                        myDBconn = pyodbc.connect(credentials.myLocalDBConn)
                        result = pd.read_sql_query("""
                            EXEC [localDB].[dbo].[SkypeGetSMS]
                                @SkypeId = '{}'
                                ,@SkypeGroup = '{}'
                                ,@SearchValue = '{}'
                                ,@Template = {}
                            """.format(message["user_id"], message["chat_id"], SearchValue, Template)
                                                   , myDBconn)

                        myDBconn.commit()
                        myDBconn.close()

                        rowCount = result['Active'].count()
                        if rowCount == 0:
                            self.SMS_Zalo_Face_Group.sendMsg(
                                'Không tìm thấy hợp đồng "{}"/ Bạn chưa có quyền sử dụng tính năng này'.format(
                                    SearchValue))
                        else:
                            if result['UserPermission'][0] == 0:
                                self.SMS_Zalo_Face_Group.sendMsg(
                                    'Không tìm thấy hợp đồng "{}"/ Bạn chưa có quyền sử dụng tính năng này'.format(
                                        SearchValue))
                            elif result['Active'][0] == 0:
                                self.SMS_Zalo_Face_Group.sendMsg(
                                    'Hợp đồng này đã tất toán')
                            # elif result['AssignedTo'][0] != 'Inhouse':
                            # elif result['AssignedTo'][0] != 'Inhouse' and result['AssignedTo'][0] != 'ManualCallJune' and result['AssignedTo'][0] != 'ManualCallJuly':
                            elif 'Inhouse' not in result['AssignedTo'][0] \
                                    and 'Manual' not in result['AssignedTo'][0] \
                                    and 'notAssign' not in result['AssignedTo'][0] \
                                    and 'Waive' not in result['AssignedTo'][0] \
                                    and result['AssignedTo'][0] != 'AiRudder' \
                                    and result['AssignedTo'][0] != 'MFI Potential' \
                                    and result['AssignedTo'][0] != 'Legal Collection' \
                                    and result['AssignedTo'][0] != 'MFI Skip Work' \
                                    and result['AssignedTo'][0] != 'Bank Skip Work':
                                self.SMS_Zalo_Face_Group.sendMsg(
                                    'Hợp đồng này hiện tại không được phân công cho Inhouse')
                            else:
                                # sent2User = self.contacts[message["user_id"]].chat
                                # sent2User.sendMsg(result['SMS'][0])

                                if Template == '1':
                                    SMS_img = Image.new('RGB', (500, 450))
                                elif Template == '2':
                                    SMS_img = Image.new('RGB', (500, 265))
                                elif Template == '3':
                                    SMS_img = Image.new('RGB', (500, 440))
                                else:
                                    SMS_img = Image.new('RGB', (500, 800))

                                SMS_draw = ImageDraw.Draw(SMS_img)

                                SMS_Wartermark = result['CrmUser'][0]
                                SMS_draw.text((20, 50), SMS_Wartermark, font=self.SMS_Wartermark_font,
                                              fill=(0, 65, 0, 30))
                                SMS_draw.text((20, 160), SMS_Wartermark, font=self.SMS_Wartermark_font,
                                              fill=(0, 65, 0, 30))
                                SMS_draw.text((20, 350), SMS_Wartermark, font=self.SMS_Wartermark_font,
                                              fill=(65, 0, 0, 30))

                                SMS_text = result['SMS'][0].replace('\r\n', '\n')
                                # print(SMS_text)
                                SMS_draw.text((10, 20), SMS_text, font=self.SMS_font)
                                # SMS_img = SMS_img.rotate(270, expand=True)

                                fullPath = os.path.join(r".\Output", SearchValue + '_SMS' + Template + '.jpg')
                                fileName = os.path.basename(fullPath)

                                SMS_img.save(fullPath)

                                try:
                                    sent2User = self.contacts[message["user_id"]].chat

                                    # sent2User = sk.contacts[msg.userId].chat
                                    # sent2User.sendMsg('test Msg')

                                    sent2User.sendFile(open(fullPath, "rb"), fileName, image=True)  # file upload


                                except ValueError:
                                    print("Oops! Check your contacts & try again...")

                                self.SMS_Zalo_Face_Group.sendMsg(
                                    'Tin nhắn mẫu đã được gửi cho bạn')

            elif message["chat_id"] == self.VID_DataUpdate_chatId:
                if self.GetData_pattern.match(message["msg"].lower()) != None:
                    SearchValue = self.LOANEID_pattern.search(message["msg"].lower()).group()
                    if SearchValue != None:
                        myDBconn = pyodbc.connect(credentials.myLocalDBConn)
                        self.VID_DataUpdate_Group.sendMsg(
                            '@{}: Yêu cầu của bạn đang được xử lí'.format(message["user_id"]))
                        result = pd.read_sql_query("""
                                                        EXEC [localDB].[dbo].[SkypeGetDataUpdate]
                                                            @SearchValue = '{}'
                                                            """.format(SearchValue), myDBconn)

                        myDBconn.commit()
                        myDBconn.close()
                        if result['data'].size > 0:
                            resultData = ''
                            for record in result['data']:
                                resultData = resultData + record.replace('\r\n', '\n')
                            # resultData = resultData.replace('"', '')
                            # fullPath = os.path.join(r".\Output", SearchValue + '_data.csv')
                            # print(fullPath)
                            # fileName = os.path.basename(fullPath)
                            # f = open(fullPath, "w+", encoding='utf-8')
                            # f.write('\ufeff')
                            # f.writelines("{}".format(resultData))
                            # f.close()

                            try:
                                sent2User = self.contacts[message["user_id"]].chat
                                sent2User.sendMsg(resultData)

                                # sent2User.sendFile(open(fullPath, "rb"), fileName, image=False)  # file upload
                                self.VID_DataUpdate_Group.sendMsg(
                                    'Thông tin hợp đồng đã được gửi cho bạn')
                            except ValueError:
                                print("Oops! Check your contacts & try again...")
                        else:
                            self.VID_DataUpdate_Group.sendMsg(
                                'Không tìm thấy thông tin hợp đồng')

            elif message["chat_id"] == self.VID_DK_chatId:
                print('PDF')
                if self.GetData_pattern.match(message["msg"].lower()) != None:
                    SearchValue = self.LOANEID_pattern.search(message["msg"].lower()).group()
                    print(SearchValue)
                    if SearchValue != None:
                        print(1)
                        myDBconn = pyodbc.connect(credentials.myLocalDBConn)
                        print(2)
                        self.VID_DK_Group.sendMsg(
                            '@{}: Yêu cầu của bạn đang được xử lí'.format(message["user_id"]))
                        result = pd.read_sql_query("""
                                                        EXEC [localDB].[dbo].[SkyGetPdf]
                                                            @userName = '{}',
                                                            @SearchValue = '{}'
                                                            """.format(message["user_id"], SearchValue), myDBconn)

                        myDBconn.commit()
                        myDBconn.close()

                        rowCount = result['Number'].count()
                        if rowCount == 0:
                            # print('Status: "{}" not found'.format(SearchValue))
                            self.VID_DK_Group.sendMsg(
                                '@{}: Không tìm thấy hợp đồng "{}"'.format(message["user_id"], SearchValue))
                        else:
                            output_folder_path = r".\PDF_Files"
                            word_file_path = ".\Template_PDF.docx"
                            template = Document(word_file_path)
                            print(3)
                            for paragraph in template.paragraphs:
                                if '{{Number}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Number}}', str(result['Number'].iloc[0]))
                                if '{{Current_day}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Current_day}}',
                                                                        str(result['Current_day'].iloc[0]))
                                if '{{Current_month}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Current_month}}',
                                                                        str(result['Current_month'].iloc[0]))
                                if '{{Client_name}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Client_name}}',
                                                                        str(result['Client_name'].iloc[0]))
                                if '{{Passport}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Passport}}', str(result['Passport'].iloc[0]))
                                if '{{Address}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Address}}', str(result['Address'].iloc[0]))
                                if '{{Ngay_vay}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Ngay_vay}}', str(result['Ngay_vay'].iloc[0]))
                                if '{{Company_name}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Company_name}}',
                                                                        str(result['Company_name'].iloc[0]))
                                if '{{DPD}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{DPD}}', str(int(result['DPD'].iloc[0])))
                                if '{{Current_day_1}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Current_day_1}}',
                                                                        str(result['Current_day_1'].iloc[0]))
                                if '{{Total_debt}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Total_debt}}',
                                                                        str(result['Total_debt'].iloc[0].replace(",",
                                                                                                                 ".")))
                                temp_word_file_path = os.path.join(output_folder_path,
                                                               f'{str(result["Number_1"].iloc[0])}.docx')
                                template.save(temp_word_file_path)

                            word_files_path = r".\PDF_Files\*.docx"
                            word_files = glob.glob(word_files_path)
                            for word_file_path in word_files:
                                # Конвертируем Word в PDF
                                convert(word_file_path, output_folder_path)

                                # Удаляем временный файл Word
                                os.remove(word_file_path)

                            fullPath = os.path.join(r".\PDF_Files", SearchValue.upper() + '.pdf')
                            fileName = os.path.basename(fullPath)

                            try:
                                sent2User = self.contacts[message["user_id"]].chat

                                # sent2User = sk.contacts[msg.userId].chat
                                # sent2User.sendMsg('test Msg')

                                sent2User.sendFile(open(fullPath, "rb"), fileName, image=False)  # file upload
                                self.VID_DK_Group.sendMsg(
                                    '@{}: Hợp đồng "{}" đã được gửi cho bạn'.format(message["user_id"], SearchValue))

                            except ValueError:
                                print("Oops! Check your contacts & try again...")
                    else:
                        self.VID_DK_Group.sendMsg(
                            'Không tìm thấy thông tin hợp đồng')

            elif message["chat_id"] == self.VID_TBKK_sunshine or message["chat_id"] == self.VID_TBKK_concentrix:
                if message["chat_id"] == self.VID_TBKK_sunshine:
                    a = self.chats[self.VID_TBKK_sunshine]
                elif message["chat_id"] == self.VID_TBKK_concentrix:
                    a = self.chats[self.VID_TBKK_concentrix]

                print('TBKK')
                if self.TBKK_pattern.match(message["msg"].lower()) != None:
                    SearchValue = self.LOANEID_pattern.search(message["msg"].lower()).group()
                    print(SearchValue)
                    if SearchValue != None:
                        myDBconn = pyodbc.connect(credentials.myLocalDBConn)
                        a.sendMsg(
                            '@{}: Yêu cầu của bạn đang được xử lí'.format(message["user_id"]))
                        result = pd.read_sql_query("""
                                                        EXEC [localDB].[dbo].[SkyGetPdf]
                                                            @userName = '{}',
                                                            @SearchValue = '{}'
                                                            """.format(message["user_id"], SearchValue), myDBconn)

                        rowCount = result['Number'].count()
                        if rowCount == 0:
                            # print('Status: "{}" not found'.format(SearchValue))
                            a.sendMsg(
                                '@{}: Không tìm thấy hợp đồng "{}"'.format(message["user_id"], SearchValue))
                        else:
                            print(result)
                            myDBconn.commit()
                            myDBconn.close()
                            output_folder_path = r".\PDF_Files_DCA"
                            word_file_path = ".\Template_PDF.docx"
                            template = Document(word_file_path)
                            print(3)
                            for paragraph in template.paragraphs:
                                if '{{Number}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Number}}', str(result['Number'].iloc[0]))
                                if '{{Current_day}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Current_day}}',
                                                                        str(result['Current_day'].iloc[0]))
                                if '{{Current_month}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Current_month}}',
                                                                        str(result['Current_month'].iloc[0]))
                                if '{{Client_name}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Client_name}}',
                                                                        str(result['Client_name'].iloc[0]))
                                if '{{Passport}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Passport}}', str(result['Passport'].iloc[0]))
                                if '{{Address}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Address}}', str(result['Address'].iloc[0]))
                                if '{{Ngay_vay}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Ngay_vay}}', str(result['Ngay_vay'].iloc[0]))
                                if '{{Company_name}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Company_name}}',
                                                                        str(result['Company_name'].iloc[0]))
                                if '{{DPD}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{DPD}}', str(int(result['DPD'].iloc[0])))
                                if '{{Current_day_1}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Current_day_1}}',
                                                                        str(result['Current_day_1'].iloc[0]))
                                if '{{Total_debt}}' in paragraph.text:
                                    paragraph.text = paragraph.text.replace('{{Total_debt}}',
                                                                        str(result['Total_debt'].iloc[0].replace(",",
                                                                                                                 ".")))
                                temp_word_file_path = os.path.join(output_folder_path,
                                                               f'{str(result["Number_1"].iloc[0])}.docx')
                                template.save(temp_word_file_path)

                            word_files_path = r".\PDF_Files_DCA\*.docx"
                            word_files = glob.glob(word_files_path)
                            for word_file_path in word_files:
                                # Конвертируем Word в PDF
                                convert(word_file_path, output_folder_path)

                                # Удаляем временный файл Word
                                os.remove(word_file_path)

                            fullPath = os.path.join(r".\PDF_Files_DCA", SearchValue.upper() + '.pdf')
                            fileName = os.path.basename(fullPath)

                            try:
                                sent2User = self.contacts[message["user_id"]].chat

                                # sent2User = sk.contacts[msg.userId].chat
                                # sent2User.sendMsg('test Msg')

                                sent2User.sendFile(open(fullPath, "rb"), fileName, image=False)  # file upload
                                a.sendMsg(
                                    '@{}: Hợp đồng "{}" đã được gửi cho bạn'.format(message["user_id"], SearchValue))

                            except ValueError:
                                print("Oops! Check your contacts & try again...")
                    else:
                        a.sendMsg(
                            'Không tìm thấy thông tin hợp đồng')

# ----START OF SCRIPT
if __name__ == "__main__":
    SkypeGetInfo = SkypeListener()
    print("Hello")
    SkypeGetInfo.loop()
